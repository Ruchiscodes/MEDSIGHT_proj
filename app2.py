# import os
# import io
# import json
# import streamlit as st
# import numpy as np
# import pydicom
# from PIL import Image as PILImage
# from docx import Document
# from docx.shared import Inches
# import fitz  # PyMuPDF
# import requests
# import google.generativeai as genai
# from pydicom.pixel_data_handlers.util import apply_modality_lut
# from dotenv import load_dotenv  # Import load_dotenv

# # ---------------- Load API Key from .env file ----------------
# load_dotenv()
# API_KEY = os.getenv("GOOGLE_API_KEY")

# # ---------------- Page Config (NEW) ----------------
# # Set page to wide layout, add title and icon
# st.set_page_config(page_title="MedSight AI Pro", page_icon="🏥", layout="wide")


# # ---------------- PubMed Tool ----------------
# def search_pubmed(query: str) -> str:
#     """Searches PubMed for top 3 results and returns as JSON string."""
#     # This print statement is now removed from the function
#     base_url = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/"
#     search_url = f"{base_url}esearch.fcgi?db=pubmed&term={query}&retmode=json&retmax=3"
#     try:
#         resp = requests.get(search_url, timeout=10)
#         resp.raise_for_status()
#         data = resp.json()
#         ids = data.get("esearchresult", {}).get("idlist", [])
#         if not ids:
#             return "No relevant articles found on PubMed."
#         ids_str = ",".join(ids)
#         summary_url = f"{base_url}esummary.fcgi?db=pubmed&id={ids_str}&retmode=json"
#         summary_resp = requests.get(summary_url, timeout=10)
#         summary_resp.raise_for_status()
#         summary_data = summary_resp.json()
#         results = []
#         for uid in ids:
#             article = summary_data["result"][uid]
#             results.append({
#                 "title": article.get("title", "N/A"),
#                 "authors": [a.get("name", "N/A") for a in article.get("authors", [])],
#                 "journal": article.get("fulljournalname", "N/A"),
#                 "pub_date": article.get("pubdate", "N/A"),
#                 "pmid": uid,
#                 "url": f"https://pubmed.ncbi.nlm.nih.gov/{uid}/", # Corrected URL
#             })
#         return json.dumps(results, indent=2)
#     except Exception as e:
#         return f"PubMed search failed: {e}"


# # ---------------- DICOM Helpers ----------------
# def get_all_dicom_metadata(dicom_file: pydicom.FileDataset) -> str:
#     metadata = "--- Full DICOM Header Metadata ---\n"
#     for tag in dicom_file.iterall():
#         if tag.keyword != "PixelData":
#             try:
#                 metadata += f"{tag.name} ({tag.tag}): {tag.value}\n"
#             except Exception:
#                 metadata += f"{tag.name} ({tag.tag}): [Unreadable]\n"
#     metadata += "---------------------------------\n"
#     return metadata


# def handle_dicom_file(uploaded_file):
#     try:
#         dicom_bytes = io.BytesIO(uploaded_file.getvalue())
#         dicom_data = pydicom.dcmread(dicom_bytes)
#         full_metadata_text = get_all_dicom_metadata(dicom_data)
#         images = []

#         try:
#             pixel_array_full = apply_modality_lut(dicom_data.pixel_array, dicom_data)
#         except Exception as e:
#             st.error(f"Error decompressing pixel data: {e}")
#             return [], None

#         if hasattr(dicom_data, "NumberOfFrames") and dicom_data.NumberOfFrames > 1:
#             st.info(f"Multi-frame DICOM detected ({dicom_data.NumberOfFrames} frames). Sampling first, middle, last.")
#             frame_indices = [0, dicom_data.NumberOfFrames // 2, dicom_data.NumberOfFrames - 1]
#             pixel_arrays = pixel_array_full
#         else:
#             frame_indices = [0]
#             pixel_arrays = [pixel_array_full]

#         for i in frame_indices:
#             pixel_array = pixel_arrays[i].astype(float)
#             # Normalize pixel array
#             if pixel_array.max() > 0:
#                 rescaled_array = (np.maximum(pixel_array, 0) / pixel_array.max()) * 255
#             else:
#                 rescaled_array = np.zeros_like(pixel_array) # Avoid division by zero
                
#             final_array = np.uint8(rescaled_array)
#             image = PILImage.fromarray(final_array)
#             if image.mode != "RGB":
#                 image = image.convert("RGB")
#             images.append(image)

#         return images, full_metadata_text

#     except Exception as e:
#         st.error(f"Error processing DICOM file: {e}")
#         return [], None


# # ---------------- PDF Helper ----------------
# def handle_pdf_file(uploaded_file):
#     try:
#         pdf_bytes = uploaded_file.getvalue()
#         doc = fitz.open(stream=pdf_bytes, filetype="pdf")
#         text_content = ""
#         images = []
#         # Extract text from all pages
#         for page in doc:
#             text_content += page.get_text() + "\n\n"
            
#         # Extract first image found in the document
#         for page_num in range(len(doc)):
#             imgs = doc.get_page_images(page_num)
#             if imgs:
#                 xref = imgs[0][0]
#                 base_img = doc.extract_image(xref)
#                 img_bytes = base_img["image"]
#                 img = PILImage.open(io.BytesIO(img_bytes))
#                 if img.mode != "RGB":
#                     img = img.convert("RGB")
#                 images.append(img)
#                 break # Only grab the first image
#         return images, text_content
#     except Exception as e:
#         st.error(f"Error processing PDF: {e}")
#         return [], None


# # ---------------- DOCX Report ----------------
# def create_docx(text, imgs):
#     doc = Document()
#     doc.add_heading("Medical Analysis Report", 0)
    
#     # Only add images if they are provided (which we set to None)
#     if imgs:
#         doc.add_heading("Analyzed Images", level=1)
#         for i, img in enumerate(imgs):
#             doc.add_paragraph(f"Image {i+1}")
#             img_io = io.BytesIO()
#             img.save(img_io, format="PNG")
#             img_io.seek(0)
#             doc.add_picture(img_io, width=Inches(5.0))
            
#     doc.add_paragraph(text)
#     return doc


# def get_docx_bytes(text, imgs):
#     doc = create_docx(text, imgs)
#     buffer = io.BytesIO()
#     doc.save(buffer)
#     buffer.seek(0)
#     return buffer.getvalue()


# # ---------------- Prompts (UPGRADED) ----------------
# # This is the section that has been changed
# NORMAL_REPORT_PROMPT = """
# You are a world-class radiology expert AI.
# Generate a **Professional Medical Report**.
# The report must be **exceptionally detailed** and clinically comprehensive for a referring physician.
# It must include:
# 1.  **Findings**: A systematic, **level-by-level analysis** of all relevant structures (e.g., vertebral bodies, intervertebral discs, thecal sac, neural foramina, facet joints, bone marrow signal).
# 2.  **Quantitative Analysis**: Specify all measurements (e.g., "5mm disc bulge"), severity (mild, moderate, severe), and anatomical location.
# 3.  **Positive and Negative Findings**: Explicitly mention both abnormal findings and relevant *negative* findings (e.g., "No evidence of cord signal abnormality," "No acute fracture").
# 4.  **Impression**: A summary of the most critical findings, listed as numbered points.
# 5.  **Differential Diagnoses**: A comprehensive list of possible diagnoses.
# 6.  **Recommendations**: Suggestions for next steps, clinical correlation, or further imaging.

# **FORMATTING**: Separate each major section (e.g., Findings, Impression) with a markdown horizontal rule (---).

# **IMPORTANT**: Do not include any signature, "Prepared by" field, or placeholder for a name/title.
# At the very end of your response, add a single line formatted exactly as:
# Search_Query: [Your 3-5 word PubMed search query here]
# """

# PATIENT_SUMMARY_PROMPT = """
# You are a medical expert with an excellent ability to communicate.
# Generate a **Patient Summary** in a narrative format.
# The report must:
# 1.  Be written at a **12th-grade reading level**.
# 2.  **Avoid complex medical terminology**. Explain findings in simple, clear terms (e.g., "The cushion between your bones is bulging" instead of "disc protrusion").
# 3.  Explain *what the findings mean* for the patient (e.g., "This bulging can press on a nerve, which is why you might feel pain in your arm.").
# 4.  **Suggestions and Advice**: Include a section with general, helpful advice, such as "What to discuss with your doctor," "Potential next steps," or "Lifestyle modifications that may help." This must be general advice, not a specific medical order.
# 5.  Maintain a professional, reassuring tone.
# 6.  **Do NOT use any emojis.**

# **FORMATTING**: Separate the main summary from the "Suggestions and Advice" section with a markdown horizontal rule (---).

# **IMPORTANT**: Do not include any signature, "Prepared by" field, or placeholder for a name/title.
# At the very end of your response, add a single line formatted exactly as:
# Search_Query: [Your 3-5 word PubMed search query here]
# """

# SPECIALIZED_REPORT_PROMPT = """
# You are an expert medical consultant specializing in medico-legal reports.
# Generate a **Specialized Report** (e.g., for an IME, Forensic, or Rebuttal analysis).
# The report must be strictly objective, formal, and data-driven, suitable for a legal or insurance review. Include:
# 1.  **Objective Findings**: A precise, quantitative, **level-by-level description** of all findings, including relevant **negative findings**. Use measurements where possible.
# 2.  **Causation Analysis**: Based *only* on the images, provide an opinion on causation (e.g., "findings are consistent with chronic degenerative change," or "findings are consistent with acute traumatic injury").
# 3.  **Apportionment (if possible)**: Attempt to differentiate pre-existing/degenerative findings from acute/traumatic findings.
# 4.  **Impairment and Prognosis**: Comment on the likely functional impairment and long-term prognosis, citing *only* the objective findings.
# 5.  **Clarity on Limitations**: Clearly state what cannot be determined from the provided files alone.

# **FORMATTING**: Separate each major section (e.g., Objective Findings, Causation) with a markdown horizontal rule (---).

# **IMPORTANT**: Do not include any signature, "Prepared by" field, or placeholder for a name/title.
# At the very end of your response, add a single line formatted exactly as:
# Search_Query: [Your 3-5 word PubMed search query here]
# """
# # ---------------- End of Changed Section ----------------


# # ---------------- Streamlit UI ----------------
# st.title("🏥 MedSight AI Pro")
# st.subheader("Advanced Medical Analysis Agent")

# # Check for API Key
# if not API_KEY:
#     st.error("🚨 GOOGLE_API_KEY not found.")
#     st.error("Please create a .env file in the same directory and add: GOOGLE_API_KEY=Your_API_Key_Here")
#     st.stop()

# genai.configure(api_key=API_KEY)


# # --- Sidebar for Controls (NEW LAYOUT) ---
# st.sidebar.title("⚙️ Controls")
# uploaded_file = st.sidebar.file_uploader("1. Upload a Medical File", type=["dcm", "dicom", "pdf","jpg", "jpeg", "png"],
# )

# report_type = st.sidebar.selectbox(
#     "2. Select Report Type",
#     ("Normal Report (Professional)", "Patient Summary", "Specialized Report (IME/Forensic/Rebuttal)")
# )

# analyze_button = st.sidebar.button("3. Analyze File", use_container_width=True, type="primary")


# # --- Main Page for Results ---
# if "analysis_result" not in st.session_state:
#     st.session_state["analysis_result"] = ""
# if "report_images" not in st.session_state:
#     st.session_state["report_images"] = []
# if "pubmed_search_keywords" not in st.session_state:
#     st.session_state["pubmed_search_keywords"] = ""

# if uploaded_file:
#     ext = uploaded_file.name.lower().split(".")[-1]
#     images, extra_text = [], ""

#     if ext in ["dcm", "dicom"]:
#         images, extra_text = handle_dicom_file(uploaded_file)
#     elif ext == "pdf":
#         images, extra_text = handle_pdf_file(uploaded_file)
#     else:
#         # Handle standard images
#         images = [PILImage.open(uploaded_file).convert("RGB")]
#         extra_text = "" # No text to extract from simple images

#     # Persist images for download button (though we don't use them in the doc)
#     st.session_state["report_images"] = images

#     # Show images in an expander
#     with st.expander("Show/Hide Uploaded Images"):
#         if images:
#             st.image(images, caption=[f"Image {i+1}" for i in range(len(images))], use_container_width=True)
#         else:
#             st.write("No images could be extracted or found in this file.")
            
#     # Show extracted text in an expander
#     if extra_text:
#         with st.expander("Show/Hide Extracted File Text/Metadata"):
#             st.text(extra_text[:3000] + "..." if len(extra_text) > 3000 else extra_text)

#     if analyze_button:
#         with st.spinner("Analyzing with Gemini model..."):
#             try:
#                 # Use a more recent model if available, otherwise stick to 1.5-flash
#                 model = genai.GenerativeModel(model_name="gemini-2.0-flash")

#                 # Trim and Resize
#                 trimmed_text = (extra_text[:3000] + " ...[trimmed]") if len(extra_text) > 3000 else extra_text
#                 resized_images = []
#                 for img in images:
#                     img_copy = img.copy()
#                     img_copy.thumbnail((512, 512))
#                     resized_images.append(img_copy)

#                 # Prepare prompt based on selection
#                 if report_type == "Normal Report (Professional)":
#                     prompt_to_send = NORMAL_REPORT_PROMPT
#                 elif report_type == "Patient Summary":
#                     prompt_to_send = PATIENT_SUMMARY_PROMPT
#                 else: # Specialized Report
#                     prompt_to_send = SPECIALIZED_REPORT_PROMPT

#                 inputs = [prompt_to_send, "--- START OF CONTEXT ---", (trimmed_text or "No text context provided."), "--- END OF CONTEXT ---"]
#                 inputs.extend(resized_images)

#                 # Token safety check
#                 token_info = model.count_tokens(inputs)
#                 if token_info.total_tokens > 1_000_000:
#                     st.error("❌ Input too large! Please upload a smaller file or fewer frames.")
#                     st.stop()

#                 # Generate
#                 response = model.generate_content(inputs)
                
#                 # --- NEW LOGIC for PubMed ---
#                 raw_response_text = response.text
#                 search_query = "radiology" # Default fallback
#                 report_text = raw_response_text # Default to full text

#                 # Reliably parse the search query and clean the report
#                 if "Search_Query:" in raw_response_text:
#                     parts = raw_response_text.split("Search_Query:")
#                     report_text = parts[0].strip() # This is the clean report
#                     search_query = parts[1].strip() # This is the keyword
                
#                 st.session_state["analysis_result"] = report_text # Store only the clean report
#                 st.session_state["pubmed_search_keywords"] = search_query # Store the keywords
#                 # --- End of New Logic ---

#                 st.markdown("### 📋 Analysis Report")
#                 st.markdown(report_text) # Display only the clean report

#             except Exception as e:
#                 st.error(f"An error occurred during analysis: {e}")


# # --- Display Download Button and PubMed Toggle (Outside the 'analyze_button' block) ---
# if st.session_state.get("analysis_result"):
#     st.download_button(
#         "⬇️ Download Report (DOCX)",
#         data=get_docx_bytes(st.session_state["analysis_result"], imgs=None), # Pass imgs=None
#         file_name=f"{report_type.split(' ')[0]}_Report.docx",
#         mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
#     )
    
#     # Toggle for PubMed References
#     if st.toggle("Show Related PubMed References", value=False):
#         with st.spinner("🔬 Searching PubMed..."):
#             # Use the new, high-quality keywords from session_state
#             pubmed_refs_json = search_pubmed(st.session_state.get("pubmed_search_keywords", "radiology"))
#             pubmed_refs_short = []
#             try:
#                 pubmed_refs_data = json.loads(pubmed_refs_json)
#                 for ref in pubmed_refs_data:
#                     title = ref.get("title", "N/A")
#                     url = ref.get("url", "")
#                     pubmed_refs_short.append(f"- [{title}]({url})")
#             except Exception:
#                 pubmed_refs_short.append(pubmed_refs_json)

#             st.markdown("### 🔬 Related PubMed References")
#             st.markdown("\n".join(pubmed_refs_short))


import os
import io
import json
import streamlit as st
import numpy as np
import pydicom
from PIL import Image as PILImage
from docx import Document
from docx.shared import Inches
import fitz  # PyMuPDF
import requests
import google.generativeai as genai
from pydicom.pixel_data_handlers.util import apply_modality_lut
from dotenv import load_dotenv  # Import load_dotenv

# ---------------- Load API Key from .env file ----------------
load_dotenv()
API_KEY = os.getenv("GOOGLE_API_KEY")

# ---------------- Page Config (NEW) ----------------
# Set page to wide layout, add title and icon
st.set_page_config(page_title="MedSight AI Pro", page_icon="🏥", layout="wide")


# ---------------- PubMed Tool ----------------
def search_pubmed(query: str) -> str:
    """Searches PubMed for top 3 results and returns as JSON string."""
    # This print statement is now removed from the function
    base_url = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/"
    search_url = f"{base_url}esearch.fcgi?db=pubmed&term={query}&retmode=json&retmax=3"
    try:
        resp = requests.get(search_url, timeout=10)
        resp.raise_for_status()
        data = resp.json()
        ids = data.get("esearchresult", {}).get("idlist", [])
        if not ids:
            return "No relevant articles found on PubMed."
        ids_str = ",".join(ids)
        summary_url = f"{base_url}esummary.fcgi?db=pubmed&id={ids_str}&retmode=json"
        summary_resp = requests.get(summary_url, timeout=10)
        summary_resp.raise_for_status()
        summary_data = summary_resp.json()
        results = []
        for uid in ids:
            article = summary_data["result"][uid]
            results.append({
                "title": article.get("title", "N/A"),
                "authors": [a.get("name", "N/A") for a in article.get("authors", [])],
                "journal": article.get("fulljournalname", "N/A"),
                "pub_date": article.get("pubdate", "N/A"),
                "pmid": uid,
                "url": f"https://pubmed.ncbi.nlm.nih.gov/{uid}/", # Corrected URL
            })
        return json.dumps(results, indent=2)
    except Exception as e:
        return f"PubMed search failed: {e}"


# ---------------- DICOM Helpers ----------------
def get_all_dicom_metadata(dicom_file: pydicom.FileDataset) -> str:
    metadata = "--- Full DICOM Header Metadata ---\n"
    for tag in dicom_file.iterall():
        if tag.keyword != "PixelData":
            try:
                metadata += f"{tag.name} ({tag.tag}): {tag.value}\n"
            except Exception:
                metadata += f"{tag.name} ({tag.tag}): [Unreadable]\n"
    metadata += "---------------------------------\n"
    return metadata


def handle_dicom_file(uploaded_file):
    try:
        dicom_bytes = io.BytesIO(uploaded_file.getvalue())
        dicom_data = pydicom.dcmread(dicom_bytes)
        full_metadata_text = get_all_dicom_metadata(dicom_data)
        images = []

        try:
            pixel_array_full = apply_modality_lut(dicom_data.pixel_array, dicom_data)
        except Exception as e:
            st.error(f"Error decompressing pixel data: {e}")
            return [], None

        if hasattr(dicom_data, "NumberOfFrames") and dicom_data.NumberOfFrames > 1:
            st.info(f"Multi-frame DICOM detected ({dicom_data.NumberOfFrames} frames). Sampling first, middle, last.")
            frame_indices = [0, dicom_data.NumberOfFrames // 2, dicom_data.NumberOfFrames - 1]
            pixel_arrays = pixel_array_full
        else:
            frame_indices = [0]
            pixel_arrays = [pixel_array_full]

        for i in frame_indices:
            pixel_array = pixel_arrays[i].astype(float)
            # Normalize pixel array
            if pixel_array.max() > 0:
                rescaled_array = (np.maximum(pixel_array, 0) / pixel_array.max()) * 255
            else:
                rescaled_array = np.zeros_like(pixel_array) # Avoid division by zero
                
            final_array = np.uint8(rescaled_array)
            image = PILImage.fromarray(final_array)
            if image.mode != "RGB":
                image = image.convert("RGB")
            images.append(image)

        return images, full_metadata_text

    except Exception as e:
        st.error(f"Error processing DICOM file: {e}")
        return [], None


# ---------------- PDF Helper ----------------
def handle_pdf_file(uploaded_file):
    try:
        pdf_bytes = uploaded_file.getvalue()
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        text_content = ""
        images = []
        # Extract text from all pages
        for page in doc:
            text_content += page.get_text() + "\n\n"
            
        # Extract first image found in the document
        for page_num in range(len(doc)):
            imgs = doc.get_page_images(page_num)
            if imgs:
                xref = imgs[0][0]
                base_img = doc.extract_image(xref)
                img_bytes = base_img["image"]
                img = PILImage.open(io.BytesIO(img_bytes))
                if img.mode != "RGB":
                    img = img.convert("RGB")
                images.append(img)
                break # Only grab the first image
        return images, text_content
    except Exception as e:
        st.error(f"Error processing PDF: {e}")
        return [], None


# ---------------- DOCX Report ----------------
def create_docx(text, imgs):
    doc = Document()
    doc.add_heading("Medical Analysis Report", 0)
    
    # Only add images if they are provided (which we set to None)
    if imgs:
        doc.add_heading("Analyzed Images", level=1)
        for i, img in enumerate(imgs):
            doc.add_paragraph(f"Image {i+1}")
            img_io = io.BytesIO()
            img.save(img_io, format="PNG")
            img_io.seek(0)
            doc.add_picture(img_io, width=Inches(5.0))
            
    doc.add_paragraph(text)
    return doc


def get_docx_bytes(text, imgs):
    doc = create_docx(text, imgs)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ---------------- Prompts (UPGRADED) ----------------
# This is the section that has been changed
NORMAL_REPORT_PROMPT = """
You are a world-class radiology expert AI.
Generate a **Professional Medical Report**.
The report must be **exceptionally detailed** and clinically comprehensive for a referring physician.
It must include:
1.  **Findings**: A systematic, **level-by-level analysis** of all relevant structures (e.g., vertebral bodies, intervertebral discs, thecal sac, neural foramina, facet joints, bone marrow signal).
2.  **Quantitative Analysis**: Specify all measurements (e.g., "5mm disc bulge"), severity (mild, moderate, severe), and anatomical location.
3.  **Positive and Negative Findings**: Explicitly mention both abnormal findings and relevant *negative* findings (e.g., "No evidence of cord signal abnormality," "No acute fracture").
4.  **Impression**: A summary of the most critical findings, listed as numbered points.
5.  **Differential Diagnoses**: A comprehensive list of possible diagnoses.
6.  **Recommendations**: Suggestions for next steps, clinical correlation, or further imaging.

**FORMATTING**: 
- Separate each major section (e.g., Findings, Impression) with a markdown horizontal rule (---).
- **Do not wrap the report in triple backticks (```).** Use standard markdown for headings and bolding.

**IMPORTANT**: Do not include any signature, "Prepared by" field, or placeholder for a name/title.
At the very end of your response, add a single line formatted exactly as:
Search_Query: [Your 3-5 word PubMed search query here]
"""

PATIENT_SUMMARY_PROMPT = """
You are a medical expert with an excellent ability to communicate.
Generate a **Patient Summary** in a narrative format.
The report must:
1.  Be written at a **12th-grade reading level**.
2.  **Avoid complex medical terminology**. Explain findings in simple, clear terms (e.g., "The cushion between your bones is bulging" instead of "disc protrusion").
3.  Explain *what the findings mean* for the patient (e.g., "This bulging can press on a nerve, which is why you might feel pain in your arm.").
4.  **Suggestions and Advice**: Include a section with general, helpful advice, such as "What to discuss with your doctor," "Potential next steps," or "Lifestyle modifications that may help." This must be general advice, not a specific medical order.
5.  Maintain a professional, reassuring tone.
6.  **Do NOT use any emojis.**

**FORMATTING**: 
- Separate the main summary from the "Suggestions and Advice" section with a markdown horizontal rule (---).
- **Do not wrap the report in triple backticks (```).** Use standard markdown for headings and bolding.

**IMPORTANT**: Do not include any signature, "Prepared by" field, or placeholder for a name/title.
At the very end of your response, add a single line formatted exactly as:
Search_Query: [Your 3-5 word PubMed search query here]
"""

SPECIALIZED_REPORT_PROMPT = """
You are an expert medical consultant specializing in medico-legal reports.
Generate a **Specialized Report** (e.g., for an IME, Forensic, or Rebuttal analysis).
The report must be strictly objective, formal, and data-driven, suitable for a legal or insurance review. Include:
1.  **Objective Findings**: A precise, quantitative, **level-by-level description** of all findings, including relevant **negative findings**. Use measurements where possible.
2.  **Causation Analysis**: Based *only* on the images, provide an opinion on causation (e.g., "findings are consistent with chronic degenerative change," or "findings are consistent with acute traumatic injury").
3.  **Apportionment (if possible)**: Attempt to differentiate pre-existing/degenerative findings from acute/traumatic findings.
4.  **Impairment and Prognosis**: Comment on the likely functional impairment and long-term prognosis, citing *only* the objective findings.
5.  **Clarity on Limitations**: Clearly state what cannot be determined from the provided files alone.

**FORMATTING**: 
- Separate each major section (e.g., Objective Findings, Causation) with a markdown horizontal rule (---).
- **Do not wrap the report in triple backticks (```).** Use standard markdown for headings and bolding.

**IMPORTANT**: Do not include any signature, "Prepared by" field, or placeholder for a name/title.
At the very end of your response, add a single line formatted exactly as:
Search_Query: [Your 3-5 word PubMed search query here]
"""
# ---------------- End of Changed Section ----------------


# ---------------- Streamlit UI ----------------
st.title("🏥 MedSight AI Pro")
st.subheader("Advanced Medical Analysis Agent")

# Check for API Key
if not API_KEY:
    st.error("🚨 GOOGLE_API_KEY not found.")
    st.error("Please create a .env file in the same directory and add: GOOGLE_API_KEY=Your_API_Key_Here")
    st.stop()

genai.configure(api_key=API_KEY)


# --- Sidebar for Controls (NEW LAYOUT) ---
st.sidebar.title("⚙️ Controls")
uploaded_file = st.sidebar.file_uploader(
    "1. Upload a Medical File",
    type=["dcm", "dicom", "pdf", "jpg", "jpeg", "png"],
)

report_type = st.sidebar.selectbox(
    "2. Select Report Type",
    ("Normal Report (Professional)", "Patient Summary", "Specialized Report (IME/Forensic/Rebuttal)")
)

analyze_button = st.sidebar.button("3. Analyze File", use_container_width=True, type="primary")


# --- Main Page for Results ---
if "analysis_result" not in st.session_state:
    st.session_state["analysis_result"] = ""
if "report_images" not in st.session_state:
    st.session_state["report_images"] = []
if "pubmed_search_keywords" not in st.session_state:
    st.session_state["pubmed_search_keywords"] = ""

if uploaded_file:
    ext = uploaded_file.name.lower().split(".")[-1]
    images, extra_text = [], ""

    if ext in ["dcm", "dicom"]:
        images, extra_text = handle_dicom_file(uploaded_file)
    elif ext == "pdf":
        images, extra_text = handle_pdf_file(uploaded_file)
    else:
        # Handle standard images
        images = [PILImage.open(uploaded_file).convert("RGB")]
        extra_text = "" # No text to extract from simple images

    # Persist images for download button (though we don't use them in the doc)
    st.session_state["report_images"] = images

    # Show images in an expander
    with st.expander("Show/Hide Uploaded Images"):
        if images:
            st.image(images, caption=[f"Image {i+1}" for i in range(len(images))], use_container_width=True)
        else:
            st.write("No images could be extracted or found in this file.")
            
    # Show extracted text in an expander
    if extra_text:
        with st.expander("Show/Hide Extracted File Text/Metadata"):
            st.text(extra_text[:3000] + "..." if len(extra_text) > 3000 else extra_text)

    if analyze_button:
        with st.spinner("Analyzing with Gemini model..."):
            try:
                # Use a more recent model if available, otherwise stick to 1.5-flash
                model = genai.GenerativeModel(model_name="gemini-pro-latest")

                # Trim and Resize
                trimmed_text = (extra_text[:3000] + " ...[trimmed]") if len(extra_text) > 3000 else extra_text
                resized_images = []
                for img in images:
                    img_copy = img.copy()
                    img_copy.thumbnail((512, 512))
                    resized_images.append(img_copy)

                # Prepare prompt based on selection
                if report_type == "Normal Report (Professional)":
                    prompt_to_send = NORMAL_REPORT_PROMPT
                elif report_type == "Patient Summary":
                    prompt_to_send = PATIENT_SUMMARY_PROMPT
                else: # Specialized Report
                    prompt_to_send = SPECIALIZED_REPORT_PROMPT

                inputs = [prompt_to_send, "--- START OF CONTEXT ---", (trimmed_text or "No text context provided."), "--- END OF CONTEXT ---"]
                inputs.extend(resized_images)

                # Token safety check
                token_info = model.count_tokens(inputs)
                if token_info.total_tokens > 1_000_000:
                    st.error("❌ Input too large! Please upload a smaller file or fewer frames.")
                    st.stop()

                # Generate
                response = model.generate_content(inputs)
                
                # --- NEW LOGIC for PubMed ---
                raw_response_text = response.text
                search_query = "radiology" # Default fallback
                report_text = raw_response_text # Default to full text

                # Reliably parse the search query and clean the report
                if "Search_Query:" in raw_response_text:
                    parts = raw_response_text.split("Search_Query:")
                    report_text = parts[0].strip() # This is the clean report
                    search_query = parts[1].strip() # This is the keyword
                
                st.session_state["analysis_result"] = report_text # Store only the clean report
                st.session_state["pubmed_search_keywords"] = search_query # Store the keywords
                # --- End of New Logic ---

                st.markdown("### 📋 Analysis Report")
                st.markdown(report_text) # Display only the clean report

            except Exception as e:
                st.error(f"An error occurred during analysis: {e}")


# --- Display Download Button and PubMed Toggle (Outside the 'analyze_button' block) ---
if st.session_state.get("analysis_result"):
    st.download_button(
        "⬇️ Download Report (DOCX)",
        data=get_docx_bytes(st.session_state["analysis_result"], imgs=None), # Pass imgs=None
        file_name=f"{report_type.split(' ')[0]}_Report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    
    # Toggle for PubMed References
    if st.toggle("Show Related PubMed References", value=False):
        with st.spinner("🔬 Searching PubMed..."):
            # Use the new, high-quality keywords from session_state
            pubmed_refs_json = search_pubmed(st.session_state.get("pubmed_search_keywords", "radiology"))
            pubmed_refs_short = []
            try:
                pubmed_refs_data = json.loads(pubmed_refs_json)
                for ref in pubmed_refs_data:
                    title = ref.get("title", "N/A")
                    url = ref.get("url", "")
                    pubmed_refs_short.append(f"- [{title}]({url})")
            except Exception:
                pubmed_refs_short.append(pubmed_refs_json)

            st.markdown("### 🔬 Related PubMed References")
            st.markdown("\n".join(pubmed_refs_short))
 